"""
Text Extraction Service (Pipeline Step 3)
Extracts raw text from PDF and DOCX files.
"""

import fitz
from docx import Document
import os


def extract_text_from_pdf(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"PDF file not found: {file_path}")

    text_parts = []
    doc = fitz.open(file_path)

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_text = page.get_text("text")
        if page_text.strip():
            text_parts.append(page_text.strip())

    doc.close()
    full_text = "\n\n".join(text_parts)

    if not full_text.strip():
        raise ValueError("No text could be extracted from the PDF. The file may be image-based.")

    return full_text


def extract_text_from_docx(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    full_text = "\n".join(text_parts)

    if not full_text.strip():
        raise ValueError("No text could be extracted from the DOCX file.")

    return full_text


def extract_text(file_path: str) -> str:
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {extension}. Only PDF and DOCX are supported.")
